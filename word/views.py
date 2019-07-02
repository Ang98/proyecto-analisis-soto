from django.shortcuts import render, HttpResponse
from django.core.files.storage import FileSystemStorage
import pandas as pd
import matplotlib.pyplot as plt
import numpy as np
from docx import Document
from renderword.settings import BASE_DIR
import os
#import docx
import json
# Create your views here.



def home(request):
    return render(request, 'word/base.html')


def cargaexcel(request):
    if request.method == 'POST':
        uploaded_file = request.FILES['document']
        fs = FileSystemStorage()
        name = fs.save(uploaded_file.name, uploaded_file)

        xls = pd.ExcelFile(BASE_DIR + fs.url(name))
        hojas = xls.sheet_names
        print(len(hojas))
        dc1 = Document()
        def mat(abcd):
            pri = 0
            val = 0
            aux = []
            for i in range(15):
                aux.append([])
                for j in range(5):
                    aux[i].append(None)

            for i in range(len(abcd)):
                sec = 0
                val = 0
                for j in range(len(abcd[i])):

                    if type(abcd[i][j]) is str:
                        val += 1
                        aux[pri][sec] = abcd[i][j]
                        sec += 1
                    else:
                        if type(abcd[i][j]) is int:
                            abcd[i][j] = str(abcd[i][j])
                            val += 1
                            aux[pri][sec] = abcd[i][j]
                            sec += 1
                        else:
                            if type(abcd[i][j]) is float:
                                if abcd[i][j]/abcd[i][j]==1:
                                    abcd[i][j] = str(abcd[i][j])
                                    val += 1
                                    aux[pri][sec] = abcd[i][j]
                                    sec += 1

                if val > 0:
                    pri += 1

            return aux
        aux=0

        for cont in range(len(hojas)-1):
            cont += 1
            df = xls.parse(hojas[cont])

            fich = df.__array__()

            real = mat(fich);
            fich = real
            print(fich)

            tit = fich[0][0]
            anio = str(fich[4][1])
            nota1 = str(fich[10][1])
            nota2 = str(fich[10][2])
            nota3 = str(fich[10][3])
            promedio = str(fich[11][1])

            dc1.add_heading(tit, 0)
            dc1.add_heading(fich[1][0], 2)
            p = dc1.add_paragraph()
            # p.add_run('\n')

            tbl1 = dc1.add_table(rows=0, cols=2)
            fila = tbl1.add_row().cells
            fila[0].text = fich[2][0]
            fila[1].text = fich[2][1]
            # p.add_run(fich[2][0] + ' \t').bold = True
            # p.add_run(fich[2][1])

            fila1 = tbl1.add_row().cells
            fila1[0].text = fich[3][0]
            fila1[1].text = fich[3][1]
            # p.add_run('\n\n')
            # p.add_run(fich[3][0] + ' \t').bold = True
            # p.add_run(fich[3][1])

            fila2 = tbl1.add_row().cells
            fila2[0].text = fich[4][0]
            fila2[1].text = anio
            # p.add_run('\n\n')
            # p.add_run(fich[4][0] + ' \t').bold = True
            # p.add_run(anio)

            dc1.add_heading(fich[5][0], 2)
            p1 = dc1.add_paragraph()
            # p1.add_run('\n')

            tbl2 = dc1.add_table(rows=0, cols=2)
            fila = tbl2.add_row().cells
            fila[0].text = fich[6][0]
            fila[1].text = fich[6][1]

            # p1.add_run(fich[6][0] + ' \t').bold = True
            # p1.add_run(fich[6][1])

            fila1 = tbl2.add_row().cells
            fila1[0].text = fich[7][0]
            fila1[1].text = fich[7][1]

            # p1.add_run('\n\n')
            # p1.add_run(fich[7][0] + ' \t').bold = True
            # p1.add_run(fich[7][1])

            dc1.add_heading(fich[8][0], 2)
            p2 = dc1.add_paragraph()
            # p2.add_run('\n')

            tbl3 = dc1.add_table(rows=0, cols=4)
            fila = tbl3.add_row().cells
            fila[0].text = fich[9][0]
            fila[1].text = fich[9][1]

            # p2.add_run(fich[9][0] + ' \t').bold = True
            # p2.add_run(fich[9][1])

            fila1 = tbl3.add_row().cells
            fila1[0].text = fich[10][0]
            fila1[1].text = nota1
            fila1[2].text = nota2
            fila1[3].text = nota3

            # p2.add_run('\n\n')
            # p2.add_run(fich[10][0] + ' \t').bold = True
            # p2.add_run(nota1+ ' \t'+nota2+ ' \t'+nota3+ ' \t')

            fila2 = tbl3.add_row().cells
            fila2[0].text = fich[11][0]
            fila2[1].text = promedio

            # p2.add_run('\n\n')
            # p2.add_run(fich[11][0] + ' \t').bold = True
            # p2.add_run(promedio)

            dc1.add_heading(fich[12][0], 2)
            p3 = dc1.add_paragraph()

            tbl4 = dc1.add_table(rows=0, cols=2)
            fila = tbl4.add_row().cells
            fila[0].text = fich[13][0]
            fila[1].text = fich[13][1]

            # p3.add_run('\n')
            # p3.add_run(fich[13][0] + ' \t').bold = True
            # p3.add_run(fich[13][1])

            fila1 = tbl4.add_row().cells
            fila1[0].text = fich[14][0]
            fila1[1].text = fich[14][1]

            # p3.add_run('\n\n')
            # p3.add_run(fich[14][0] + ' \t').bold = True
            # p3.add_run(fich[14][1])

            # dc1.add_picture('C:/Users/Angel/Desktop/prueba/excel/imgplot.png')



            """
            nombimg = 'img' + str(aux) + '.jpg'
            num = int(nota)
            nomb = str(fich[2][1])
            nombre = ("", nomb, "")
            posicion_y = np.arange(3)
            unidad = (0, num, 0)

            plt.barh(posicion_y, unidad, align="center")
            plt.yticks(posicion_y, nombre)
            plt.xlabel("NOTA")
            plt.title("NOTAS")
            plt.savefig(BASE_DIR+'/media/'+nombimg)



            dc1.add_picture(BASE_DIR+'/media/'+nombimg)

            os.remove(BASE_DIR+'/media/'+nombimg)
            """

            dc1.add_page_break()


        nombre_archivo = "Proyecto.docx"
        response = HttpResponse(content_type="application/msword")
        contenido = "attachment; filename= {0}".format(nombre_archivo)
        response["Content-Disposition"] = contenido
        dc1.save(response)

        return response
        #return render(request, 'word/index.html', {'url':fs.url(name)})

    return render(request, 'word/index.html')

